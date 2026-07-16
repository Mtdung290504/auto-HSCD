"""
word_editor.py - Universal Word document editor
Part of the word_editor skill.

Usage:
  python word_editor.py --read  <file.docx>       # Đọc cấu trúc file
  python word_editor.py --apply <changes.json>    # Thực thi thay thế

Lưu ý kỹ thuật:
- --read dùng python-docx (chỉ đọc body paragraphs, không tính paragraph trong bảng)
- --apply dùng Word COM — replace trong Range cụ thể, không phải toàn file
- Không hỗ trợ paragraph_index vì index python-docx và COM không tương đương khi có bảng
"""

import sys
import os
import json

sys.stdout.reconfigure(encoding="utf-8")

# Giới hạn của Word Find/Replace engine
WORD_REPLACE_MAX_LEN = 250  # Thực tế giới hạn ~255, để an toàn dùng 250


# ===========================================================
# MODE: --read
# Dùng python-docx để đọc cấu trúc (nhanh, không cần COM).
# CHÚ Ý: paragraphs chỉ liệt kê body-level paragraphs,
# KHÔNG bao gồm paragraph bên trong ô bảng.
# Dùng kết quả này để xác định anchor và find text,
# KHÔNG dùng index paragraph để targeting trong --apply.
# ===========================================================
def cmd_read(file_path):
    import docx

    doc = docx.Document(file_path)
    result = {
        "file": file_path,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "warning": "paragraph index trong output này KHÔNG tương đương với COM index khi có bảng. Chỉ dùng để xác định anchor/find text.",
        "paragraphs": [],
        "tables": [],
    }

    for idx, para in enumerate(doc.paragraphs):
        result["paragraphs"].append({"index": idx, "text": para.text.strip()})

    for t_idx, table in enumerate(doc.tables):
        table_data = {
            "index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cells": [],
        }
        seen = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id in seen:
                    continue
                seen.add(cell_id)
                table_data["cells"].append(
                    {
                        "row": r_idx,
                        "col": c_idx,
                        "text": cell.text.strip().replace("\n", " "),
                        "note": (
                            "Merged cell: tọa độ COM có thể lệch"
                            if _is_merged(cell)
                            else None
                        ),
                    }
                )
        result["tables"].append(table_data)

    print(json.dumps(result, ensure_ascii=False, indent=2))


def _is_merged(cell):
    """Kiểm tra ô bảng có phải merge cell không."""
    try:
        tc = cell._tc
        return (
            tc.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge"
            )
            is not None
            or tc.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan"
            )
            is not None
        )
    except Exception:
        return False


# ===========================================================
# MODE: --apply
# Dùng Word COM để thực thi thay thế trong Range cụ thể.
# ===========================================================
def cmd_apply(changes_file):
    import win32com.client
    import pythoncom

    with open(changes_file, "r", encoding="utf-8") as f:
        changes = json.load(f)

    target_file = os.path.abspath(changes["target_file"])
    replacements = changes.get("replacements", [])
    show_ui = changes.get("show_ui", False)

    if not os.path.exists(target_file):
        print(f"[ERROR] File không tồn tại: {target_file}")
        sys.exit(1)

    if not replacements:
        print("[WARN] Không có replacement nào trong changes.json.")
        return

    pythoncom.CoInitialize()
    word = None
    doc = None
    saved = False

    try:
        print("[INFO] Khởi động Microsoft Word...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = show_ui
        if show_ui:
            word.WindowState = 1  # wdWindowStateNormal

        print(f"[INFO] Mở file: {target_file}")
        doc = word.Documents.Open(target_file)

        total_replaced = 0

        for i, rule in enumerate(replacements):
            scope = rule.get("scope", "paragraph")
            find_text = rule.get("find", "")
            replace_text = rule.get("replace", "")
            # max_replacements: số lần tối đa cần thay (0 = không giới hạn, thay tất cả trong Range)
            max_replacements = rule.get("max_replacements", 1)
            comment = rule.get("comment", f"Rule {i}")

            if not find_text:
                print(f"  [SKIP] Rule {i}: 'find' rỗng. | {comment}")
                continue

            if "^" in find_text or "^" in replace_text:
                print(
                    f"  [WARN] Rule {i}: Chuỗi chứa '^' — Word sẽ diễn giải như ký tự đặc biệt. Kiểm tra kết quả. | {comment}"
                )

            target_range = None

            # --------------------------------------------------
            # Scope: table_cell — chính xác nhất
            # --------------------------------------------------
            if scope == "table_cell":
                t_idx = rule.get("table_index", 0) + 1  # 0-based → 1-based
                r_idx = rule.get("row_index", 0) + 1
                c_idx = rule.get("col_index", 0) + 1
                try:
                    target_range = doc.Tables(t_idx).Cell(r_idx, c_idx).Range
                except Exception as e:
                    print(
                        f"  [ERROR] Rule {i}: Không lấy được Table[{t_idx-1}] R{r_idx-1}C{c_idx-1}: {e} | {comment}"
                    )
                    continue

            # --------------------------------------------------
            # Scope: paragraph — chỉ hỗ trợ anchor (không index)
            # --------------------------------------------------
            elif scope == "paragraph":
                anchor = rule.get("anchor")
                if not anchor:
                    print(
                        f"  [ERROR] Rule {i}: Scope 'paragraph' yêu cầu 'anchor'. paragraph_index không được hỗ trợ. | {comment}"
                    )
                    continue

                found_para = None
                for para in doc.Paragraphs:
                    para_text = para.Range.Text
                    if anchor in para_text and find_text in para_text:
                        found_para = para
                        break

                if found_para is None:
                    print(
                        f"  [MISS] Rule {i}: Không tìm thấy anchor='{anchor[:50]}' | {comment}"
                    )
                    continue

                target_range = found_para.Range

            else:
                print(
                    f"  [ERROR] Rule {i}: Scope không hợp lệ '{scope}'. Dùng 'paragraph' hoặc 'table_cell'. | {comment}"
                )
                continue

            # --------------------------------------------------
            # Thực thi replacement trong Range
            # --------------------------------------------------
            n = _replace_in_range(
                word_range=target_range,
                find_text=find_text,
                replace_text=replace_text,
                max_replacements=max_replacements,
                word_app=word,
            )

            if n > 0:
                limit_str = "ALL" if max_replacements == 0 else str(max_replacements)
                print(
                    f"  [OK]   Rule {i}: [max={limit_str}] '{find_text}' → '{replace_text}' ({n}x thực tế) | {comment}"
                )
                total_replaced += n
            else:
                print(
                    f"  [MISS] Rule {i}: '{find_text}' không tìm thấy trong target range | {comment}"
                )

        print(
            f"\n[INFO] Tổng replacements: {total_replaced}/{len(replacements)} rules khớp"
        )
        print("[INFO] Đang lưu file...")
        doc.Save()
        saved = True
        print(f"[SUCCESS] Đã lưu: {target_file}")

    except Exception as e:
        print(f"[FATAL] {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)

    finally:
        if doc is not None:
            if show_ui and not saved:
                # Lỗi trước khi save, giữ Word mở để debug
                print("[INFO] Lỗi trước khi save — giữ Word mở để kiểm tra.")
            elif show_ui and saved:
                # Lưu thành công, giữ Word mở theo yêu cầu
                print("[INFO] Word UI vẫn mở. Đóng thủ công khi xong.")
            else:
                doc.Close(False)
        if word is not None and not show_ui:
            word.Quit()
        pythoncom.CoUninitialize()


# ===========================================================
# HELPER: Thực thi replace trong một Range cụ thể
#
# Dùng loop Find.Execute() + Selection.TypeText() cho MỌI trường hợp.
# Lý do không dùng Find.Execute(Replace=wdReplaceOne/All):
#   - wdReplaceOne/All chỉ có 2 trạng thái, không hỗ trợ N lần
#   - Trả về bool, không trả được count thực tế
#   - Bị giới hạn 255 ký tự với replace_text dài
#
# Loop TypeText giải quyết cả 3 vấn đề:
#   - Hỗ trợ chính xác max_replacements lần
#   - Trả về count thực tế
#   - Không giới hạn độ dài replace_text
#   - Formatting gốc được bảo toàn (TypeText thay thế Selection hiện tại)
#
# max_replacements: số lần tối đa (0 = không giới hạn)
# ===========================================================
def _replace_in_range(
    word_range, find_text, replace_text, max_replacements=1, word_app=None
):
    if len(replace_text) > WORD_REPLACE_MAX_LEN:
        print(
            f"    [INFO] replace_text dài ({len(replace_text)} ký tự) — vượt giới hạn 255 ký tự của Word Find engine, dùng TypeText."
        )

    find_obj = word_range.Find
    find_obj.ClearFormatting()
    find_obj.Text = find_text
    find_obj.Forward = True
    find_obj.Wrap = 0  # wdFindStop — không thoát khỏi Range
    find_obj.MatchCase = True
    find_obj.MatchWholeWord = False
    find_obj.MatchWildcards = False

    count = 0
    unlimited = max_replacements == 0

    while find_obj.Execute():
        # Find.Execute() đặt Selection vào vùng text vừa tìm được
        # TypeText thay thế Selection đó — Word tự xử lý formatting
        word_app.Selection.TypeText(replace_text)
        count += 1
        if not unlimited and count >= max_replacements:
            break

    return count


# ===========================================================
# ENTRY POINT
# ===========================================================
def main():
    args = sys.argv[1:]

    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    mode = args[0]

    if mode == "--read":
        if len(args) < 2:
            print("[ERROR] --read cần đường dẫn file .docx")
            sys.exit(1)
        file_path = os.path.abspath(args[1])
        if not os.path.exists(file_path):
            print(f"[ERROR] File không tồn tại: {file_path}")
            sys.exit(1)
        cmd_read(file_path)

    elif mode == "--apply":
        if len(args) < 2:
            print("[ERROR] --apply cần đường dẫn changes.json")
            sys.exit(1)
        changes_file = os.path.abspath(args[1])
        if not os.path.exists(changes_file):
            print(f"[ERROR] File không tồn tại: {changes_file}")
            sys.exit(1)
        cmd_apply(changes_file)

    else:
        print(f"[ERROR] Mode không hợp lệ '{mode}'. Dùng --read hoặc --apply.")
        sys.exit(1)


if __name__ == "__main__":
    main()
